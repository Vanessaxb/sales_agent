from langchain_tavily import TavilySearch
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain.docstore.document import Document as LC_Document
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
import streamlit as st
from dotenv import load_dotenv

# Removing the Markdowns
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Adding clickable links in exported files
import re
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE as RT



# Optional document parsing libraries
import PyPDF2
import docx
from io import BytesIO
import io
import fpdf
from fpdf import FPDF
import os


#print("====================================================================**********************")
#print(f"fpdf2 version: {fpdf.__version__}")

# Check DejaVu font path
font_path = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")

if os.path.exists(font_path):
    print(f"DejaVuSans.ttf found at: {font_path}")
else:
    print(f"WARNING: DejaVuSans.ttf NOT found at: {font_path}. Unicode characters may fail in PDF.")
#print("====================================================================")


# Load environment variables
from dotenv import load_dotenv, dotenv_values

load_dotenv(
    dotenv_path="/Users/vanessabeattie/Documents/Perscholas/AIPE-03/sales_agent/.env",
    override=True,
)


#print("Using GROQ_API_KEY (prefix):", os.environ["GROQ_API_KEY"][:8])

# Initialize Groq Model
llm = ChatGroq(model="openai/gpt-oss-20b", api_key=os.getenv("GROQ_API_KEY"))
#print(os.getenv("GROQ_API_KEY"))

# Initialize SearchTool
search_tool = TavilySearch(topic="general", max_results=2)

# Generate Insights
def generate_insights(company_name, product_name, company_url, company_competitors, product_doc_text=None, value_proposition=None, target_customer=None, product_category="Not Specified",
                       detailed=True):
    
    # Handle missing inputs gracefully
    if value_proposition is None:
        value_proposition = "Not provided"
    if target_customer is None:
        target_customer = "Not specified"

    # Perform the search
    search_query = f"Site:{company_url} company strategy, leadership, competitors, business model"
    search_results = search_tool.invoke(search_query)
    print("Search Results: ", search_results)

    # Prepare document text for the LLM
    if product_doc_text:
        doc_text_for_prompt = f"""
The following is the full content of the uploaded product document. 
Analyze it carefully and include any insights in your report:

{product_doc_text}
"""
    else:
        doc_text_for_prompt = "No product document was uploaded."
    
    # Build Prompt
    if detailed:
        prompt_content = f"""
Step 1: Analyze the company information below carefully.
{search_results}

Step 2: Focus on the company, product, and competitors:
- Company: {company_name}
- Product: {product_name}
- Product Category: {product_category}
- Value Proposition: {value_proposition}
- Target Customer: {target_customer}
- Competitors: {company_competitors}
- Product Document Extract: {product_doc_text if product_doc_text else "Not Provided"}

Step 3: Generate a one-page report including:
1. Company strategy related to {product_name}.
2. Product positioning and category fit ({product_category}).   
3. Value proposition alignment with company needs.
4. Possible competitors or partnerships (including {company_competitors}).
5. Leadership and decision-makers relevant to this area.
6. Product/Strategy Summary: For public companies, include insights from 10-Ks, annual reports, or other relevant documents available online.
7. Article Links: Provide links to full articles, press releases, or other sources mentioned in the output.
8. Insights extracted from the uploaded product document (if available).

Step 4: Format the output in clear sections with bullet points, headings, and concise language.

Step 5: If some information is missing, indicate "Not Found" instead of making it up.
"""
    else:
        prompt_content = f"""
Step 1: Analyze the company information below carefully.
Analyze the company info below:
{search_results}

Step 2: Focus on the company, product, and competitors:
Company: {company_name}
Product: {product_name}
Product Category: {product_category}
Value Proposition: {value_proposition}
Target Customer: {target_customer}
Competitors: {company_competitors}
Product Document: {product_doc_text if product_doc_text else 'Not Provided'}

Step 3: Generate a concise summary with max 300 words. 
Step 4: Format the output in clear sections in table format, bold headings, and concise language.
Step 5: If some information is missing, indicate "Not Found" instead of making it up.
"""
    
    messages = [
        SystemMessage("You are a sales assistant that provides concise and structured insights."),
        HumanMessage(content=prompt_content)
    ]

    model_response = llm.invoke(messages)
    print("\nModel Response: ", model_response.content)
    return model_response.content

# ======= Helper for clickable links in DOCX =======
def add_hyperlink(paragraph, text, url):
    """
    Insert a real clickable hyperlink into a python-docx paragraph.
    Uses an external relationship and applies the built-in 'Hyperlink' style (blue + underlined).
    """
    # Create relationship id to external target
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
     
     # Build <w:hyperlink r:id="...">
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Build the run that will display the text
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

       # Apply built-in Hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Force blue + underline in case style is missing
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # blue
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)


    new_run.append(rPr)

    # Text node (preserve spaces just in case)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# ====== Functions for Export Using LangChain Document ======
def create_lc_document(report_text, company_name, product_name):
    return LC_Document(
        page_content=report_text,
        metadata={"company": company_name, "product": product_name}
    )

def export_docx(lc_doc):
    doc = docx.Document()
    doc.add_heading(f"{lc_doc.metadata['company']} - {lc_doc.metadata['product']} Report", 0)

    lines = [ln.rstrip() for ln in lc_doc.page_content.splitlines()]
    url_pattern = r'(https?://[^\s\)\]]+)'  # avoid grabbing trailing ) ]
    in_links_section = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # blank line
        if not line:
            doc.add_paragraph("")
            i += 1
            continue

        # detect the "Article Links & Press Releases" header and make it bold
        if re.match(r'(?i)^article\s+links?.*press\s*releases?', line):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            in_links_section = True
            i += 1
            continue

        # --- Special handling inside the links section: bullet title + next-line URL
        if in_links_section:
            # bullet title on this line?
            m = re.match(r'^[\-\u2022•]\s*(.+)$', line)  # -, •, bullet char
            next_line = lines[i+1].strip() if (i + 1) < len(lines) else ""

            if m and re.search(url_pattern, next_line):
                title_text = m.group(1)
                url = re.search(url_pattern, next_line).group(1)

                # create a bullet paragraph whose text is a hyperlink to the URL
                para = doc.add_paragraph(style='ListBullet')  # keep your bullet look
                add_hyperlink(para, title_text, url)

                i += 2  # consume title + url line
                continue

            # if the line itself has inline URL(s), link them in-place
            urls_here = re.findall(url_pattern, line)
            if urls_here:
                para = doc.add_paragraph(style='ListBullet' if re.match(r'^[\-\u2022•]\s*', line) else None)
                parts = re.split(url_pattern, line)
                for part in parts:
                    if re.match(url_pattern, part or ""):
                        add_hyperlink(para, part, part)
                    else:
                        para.add_run(part)
                i += 1
                continue

        # --- Outside (or after) the links section: general formatting you already had

        # Inline URLs anywhere → make them clickable
        urls_here = re.findall(url_pattern, line)
        if urls_here:
            para = doc.add_paragraph()
            parts = re.split(url_pattern, line)
            for part in parts:
                if re.match(url_pattern, part or ""):
                    add_hyperlink(para, part, part)
                else:
                    para.add_run(part)
            i += 1
            continue

        # Bullets
        if line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style='ListBullet')
            i += 1
            continue

        # Headings (your rule)
        if line.endswith(":") or line.startswith("Step") or line[0].isdigit():
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            i += 1
            continue

        # Normal paragraph
        doc.add_paragraph(line)
        i += 1

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ====== PDF Export Function ==============================

class PDF(FPDF):
    def header(self):
        self.set_font("DejaVu", 'B', 16)
        self.cell(0, 10, "Sales Insights Report - Apple", ln=True, align='C')
        self.ln(5)  # Line break

########def export_pdf(report_content: dict, company_name: str):
def export_pdf(report_content, company_name: str):
    pdf = FPDF('P','mm','letter')
    pdf.set_auto_page_break(auto=True, margin=10)  # bottom margin
    pdf.set_left_margin(10)   # left margin
    pdf.set_right_margin(10)  # right margin
    pdf.add_page()

     # ----- Load monospace font here -----
    monospace_font_path = os.path.join(os.path.dirname(__file__), "DejaVuSansMono.ttf")
    if os.path.exists(monospace_font_path):
        pdf.add_font('DejaVuMono', '', monospace_font_path, uni=True)
        pdf.add_font('DejaVuMono', 'B', monospace_font_path, uni=True)
    else:
        #print("WARNING: DejaVuSansMono.ttf not found. Falling back to Courier.")
        pdf.set_font('Courier', '', 12)

    # Title
    if os.path.exists(monospace_font_path):
        #pdf.add_font('DejaVuMono', 'B', monospace_font_path, uni=True)
        pdf.set_font('DejaVuMono', 'B', 18)
    else:
        pdf.set_font('Courier', 'B', 18)

    pdf.cell(0, 10, f'Sales Insights Report - {company_name}', 0, 1, 'C')
    pdf.ln(10)

    # Content
    if isinstance(report_content, dict):
        content_text = "\n".join(f"{k}: {v}" for k, v in report_content.items())
    else:
        content_text = str(report_content)

    
    
    # Calculate page width (respect margins)
    page_width = pdf.w - pdf.l_margin - pdf.r_margin
    import re

   # Convert dict to string if needed
    if isinstance(report_content, dict):
        content_text = "\n".join(f"{k}: {v}" for k, v in report_content.items())
    else:
        content_text = str(report_content)

    for line in content_text.split('\n'):
        # Remove all leading whitespace
        line = re.sub(r'^[\s\u00A0\t]+', '', line)

        if not line:
            pdf.ln(3)
            continue

        # Detect headings: lines ending with ":" or starting with numbers
        is_heading = line.endswith(":") or re.match(r'^\d+[\.\)]', line)

        if is_heading:
            # Bold and bigger font for heading
            if os.path.exists(monospace_font_path):
                pdf.set_font('DejaVuMono', 'B', 16)
            else:
                pdf.set_font('Courier', 'B', 16)
        else:
            # Normal font for content
            if os.path.exists(monospace_font_path):
                pdf.set_font('DejaVuMono', '', 12)
            else:
                pdf.set_font('Courier', '', 12)

         # Detect URLs
        urls = re.findall(r'(https?://\S+)', line)
        if urls:
            parts = re.split(r'(https?://\S+)', line)
            for part in parts:
                if re.match(r'https?://\S+', part):
                    pdf.set_text_color(0, 0, 255)
                    pdf.set_font('', 'U')
                    pdf.write(6, part, link=part)
                    pdf.set_font('', '')
                    pdf.set_text_color(0, 0, 0)
                else:
                    pdf.write(6, part)
            pdf.ln(6)
        else:     
        # Normalize bullets/dashes
            if line.startswith("-"):
                line = "• " + line[1:].lstrip()

            # Reset x-position for each line
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(page_width, 6, line)


    # Output PDF to BytesIO
    pdf_output = io.BytesIO()
    pdf_output.write(pdf.output(dest='S').encode('latin1'))  # or 'utf-8' if using unicode fonts
    pdf_output.write(pdf.output(dest='S'))  # works if output is already bytes

    #pdf_output.write(pdf.output(dest='S'))  # write bytearray directly
    pdf_output.seek(0)
    return pdf_output




# ============= Streamlit UI ==============
st.title("Sales Agent")
st.subheader("Generate a Sales Report")
st.divider()

# Inputs
# company name
company_name = st.text_input("Company Name")
# company URL
company_url = st.text_input("Company URL")
# product name
product_name = st.text_input("Product Name")
# product category
product_category = st.text_input("Product Category (or leave blank to infer)")
# company competitors
company_competitors = st.text_input("Company Competitors")
# Value Proposition
value_proposition = st.text_area("Value Proposition (one sentence)")
# Target Customer
target_customer = st.text_input("Target Customer (decision-maker name/title)")


# Optional product document upload
uploaded_file = st.file_uploader("Upload Product Overview Sheet/Deck (PDF/DOCX optional)", type=["pdf", "docx"])
product_doc_text = None

if uploaded_file is not None:
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()  # read the file once
    if filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        product_doc_text = text.strip() or None  # assign None if empty
        
    elif filename.endswith(".docx"):
        doc = docx.Document(BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        product_doc_text = text.strip() or None  # assign None if empty
        #print("=================================================================")
        print(product_doc_text)
    else:
        st.warning("Unsupported file type. Please upload PDF or DOCX.")

# If still None, keep it as None; generate_insights will handle "Not Provided"

# detailed report
detailed_report = st.checkbox("Generate Detailed Report", value=True)

if st.button("Generate Report"):
    if company_name and company_url:
        with st.spinner("Generating Report..."):
            result = generate_insights(company_name, product_name, company_url, company_competitors, product_doc_text, value_proposition=value_proposition,
                target_customer=target_customer, product_category=product_category, detailed=detailed_report)

            st.divider()
            st.write(result)

            # Create LangChain Document
            lc_doc = create_lc_document(result, company_name, product_name)

            # Export Options
            st.subheader("Export Report")
            docx_stream = export_docx(lc_doc)
            #pdf_stream = export_pdf(lc_doc.page_content, lc_doc.metadata['company'])
            pdf_prompt = PromptTemplate.from_template("Generate a pdf friendly version of the following content in plain text. Don't use any markdown language.Do NOT indent bullets or use tabs or extra spaces. Use '-' only at the start of the line with a single space after it. :\n{content}")
            pdf_chain = pdf_prompt | llm | StrOutputParser()
            #pdf_stream = export_pdf(pdf_chain.invoke({"content": lc_doc.page_content}), lc_doc.metadata['company'])
            formated_pdf = pdf_chain.invoke({"content": lc_doc.page_content})
            print("******************************+++++++++++++++++++++++++++++++formated_pdf\n", formated_pdf)
            pdf_stream = export_pdf(formated_pdf, lc_doc.metadata['company'])

            st.download_button(
                label="Download DOCX",
                data=docx_stream,
                file_name=f"{company_name}_{product_name}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.download_button(
                label="Download PDF",
                data=pdf_stream,
                file_name=f"{company_name}_{product_name}_report.pdf",
                mime="application/pdf"
            )
    else:
        st.warning("Please enter a company name and URL")

